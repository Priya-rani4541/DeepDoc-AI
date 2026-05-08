import os
import pypdf
import docx
from config import DATA_PATH


class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}


# =========================
# TEXT CLEANING FUNCTION
# =========================
def clean_text(text):
    if not text:
        return ""

    # Remove extra spaces/newlines
    text = text.replace("\n", " ")
    text = " ".join(text.split())

    return text.strip()


# =========================
# NOISE FILTERING FUNCTION
# =========================
def is_valid_text(text):
    if not text:
        return False

    text_lower = text.lower()

    # Skip very short content
    if len(text) < 50:
        return False

    # Remove noisy/index pages
    blocked_words = [
        "table of contents",
        "contents",
        "index",
        "chapter list",
    ]

    if any(word in text_lower for word in blocked_words):
        return False

    return True


# =========================
# LOAD PDF
# =========================
def load_pdf(path):
    documents = []

    try:
        reader = pypdf.PdfReader(path)

        for i, page in enumerate(reader.pages):
            text = page.extract_text()

            text = clean_text(text)

            if not is_valid_text(text):
                continue

            documents.append(
                Document(
                    text,
                    {
                        "source": os.path.basename(path),
                        "page": i + 1
                    }
                )
            )

    except Exception as e:
        print(f"Error loading PDF {path}: {e}")

    return documents


# =========================
# LOAD TXT
# =========================
def load_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        text = clean_text(text)

        if not is_valid_text(text):
            return []

        return [
            Document(
                text,
                {
                    "source": os.path.basename(path),
                    "page": 1
                }
            )
        ]

    except Exception as e:
        print(f"Error loading TXT {path}: {e}")
        return []


# =========================
# LOAD DOCX
# =========================
def load_docx(path):
    try:
        doc = docx.Document(path)

        text = "\n".join([para.text for para in doc.paragraphs])

        text = clean_text(text)

        if not is_valid_text(text):
            return []

        return [
            Document(
                text,
                {
                    "source": os.path.basename(path),
                    "page": 1
                }
            )
        ]

    except Exception as e:
        print(f"Error loading DOCX {path}: {e}")
        return []


# =========================
# MAIN DOCUMENT LOADER
# =========================
def load_documents():
    all_docs = []

    if not os.path.exists(DATA_PATH):
        print(f"Data path {DATA_PATH} does not exist.")
        return []

    for file in os.listdir(DATA_PATH):

        path = os.path.join(DATA_PATH, file)

        if file.endswith(".pdf"):
            all_docs.extend(load_pdf(path))

        elif file.endswith(".txt"):
            all_docs.extend(load_txt(path))

        elif file.endswith(".docx"):
            all_docs.extend(load_docx(path))

    return all_docs